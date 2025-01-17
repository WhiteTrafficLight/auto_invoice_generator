import os
import sqlite3
from datetime import datetime
from docx import Document

class InvoiceGenerator:
    def __init__(self, template_path, output_dir, db_path="customers.db"):
        self.template_path = template_path
        self.output_dir = output_dir
        self.db_path = db_path
        self.price_item1 = 3.40  # € per unit
        self.price_item2 = 3.60  # € per unit
        self.tax_rate = 0.07    # 7%
        self.initialize_database()

    def initialize_database(self):
        """
        Initialize the SQLite database for storing customer information.
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS customers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE NOT NULL
            )
        """)
        conn.commit()
        conn.close()

    def add_customer(self, customer_name):
        """
        Add a new customer to the database.
        Args:
            customer_name (str): The name of the customer to add.
        Returns:
            bool: True if the customer was added, False if the name already exists.
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        try:
            cursor.execute("INSERT INTO customers (name) VALUES (?)", (customer_name,))
            conn.commit()
            conn.close()
            return True
        except sqlite3.IntegrityError:
            conn.close()
            return False

    def get_customers(self):
        """
        Retrieve the list of all customers from the database.
        Returns:
            list: List of customer names.
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM customers")
        customers = [row[0] for row in cursor.fetchall()]
        conn.close()
        return customers

    # Other methods remain unchanged


    def format_decimal(self, value):
        """
        Format a decimal value using a comma as the decimal point (German formatting).
        Args:
            value (float): The value to format.
        Returns:
            str: Formatted string with a comma as the decimal separator.
        """
        return f"{value:,.2f}".replace('.', ',')

    def calculate_totals(self, quantity_item1, quantity_item2):
        """
        Calculate total prices and tax based on quantities.
        Args:
            quantity_item1 (int): Quantity of item1.
            quantity_item2 (int): Quantity of item2.
        Returns:
            dict: Calculated totals including tax.
        """
        total_price_item1 = quantity_item1 * self.price_item1
        total_price_item2 = quantity_item2 * self.price_item2
        total_price = total_price_item1 + total_price_item2
        tax = total_price * self.tax_rate
        total_price_with_tax = total_price + tax

        return {
            "total_price_item1": total_price_item1,
            "total_price_item2": total_price_item2,
            "total_price": total_price,
            "tax": tax,
            "total_price_with_tax": total_price_with_tax,
        }

    def replace_placeholders(self, doc, placeholders):
        """
        Replace placeholders in a Word document, including text in paragraphs and tables.
        Args:
            doc (Document): The Word document object.
            placeholders (dict): A dictionary of placeholders and their replacement values.
        """
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in placeholders.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{{{key}}}}}", value)

    def generate_invoice_number(self, year_dir):
        """
        Generate an invoice number based on the number of files in the corresponding year directory.
        Args:
            year_dir (str): Path to the year directory.
        Returns:
            int: The generated invoice number.
        """
        if not os.path.exists(year_dir):
            return 1
        files = [f for f in os.listdir(year_dir) if f.endswith('.docx')]
        return len(files) + 1
    
    def get_week_of_month(self, date):
        """
        Calculate the week of the month for a given date.
        Args:
            date (datetime): The date to calculate the week for.
        Returns:
            int: The week of the month.
        """
        first_day_of_month = date.replace(day=1)
        return (date.day + first_day_of_month.weekday()) // 7 + 1

    def generate_invoice(self, user_data):
        """
        Generate an invoice and save it in the correct directory.
        Args:
            user_data (dict): Input data containing 'customer', 'year', 'month', 'day', 'quantity_item1', 'quantity_item2'.
        Returns:
            str: Path to the saved invoice file.
        """
        # Load the Word template
        doc = Document(self.template_path)
        
        # Get today's date
        today = datetime.today()
        week_of_month = self.get_week_of_month(today)
        
        # Determine directory structure based on today's date
        year_dir = os.path.join(self.output_dir, str(today.year))
        month_dir = os.path.join(year_dir, f"{today.year}-{today.month:02d}")
        week_dir = os.path.join(month_dir, f"{week_of_month:02d}")
        os.makedirs(week_dir, exist_ok=True)

        # Generate invoice number
        invoice_number = self.generate_invoice_number(year_dir)

        # Generate invoice number
        invoice_number = self.generate_invoice_number(year_dir)

        # Extract quantities and calculate totals
        quantity_item1 = int(user_data["quantity_item1"])
        quantity_item2 = int(user_data["quantity_item2"])
        totals = self.calculate_totals(quantity_item1, quantity_item2)

        # Combine user data and calculated totals
        formatted_date = f"{user_data['day']:02d}-{user_data['month']:02d}-{user_data['year']}"
        placeholders = {
            **user_data,
            "order_date": formatted_date,
            "invoice_nr": f"{user_data['year']}-{invoice_number:03d}",
            **{key: self.format_decimal(value) for key, value in totals.items()}
        }

        # Replace placeholders in the document
        self.replace_placeholders(doc, placeholders)

        # Save the invoice
        invoice_name = f"{user_data['customer']}-{invoice_number:03d}.docx"
        invoice_path = os.path.join(week_dir, invoice_name)
        doc.save(invoice_path)

        return invoice_path




